#!/usr/bin/env python3
"""Gera o DOCX desta atividade a partir de report.json (padrao PDI senior)."""
import json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
with open(os.path.join(HERE, "report.json"), encoding="utf-8") as f:
    D = json.load(f)

doc = Document()
st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(3); s.right_margin = Cm(3)

def cover():
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_before = Pt(110)
    r = p.add_run('PDI'); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x8B,0x45,0x13)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(D['title']); r.bold = True; r.font.size = Pt(23)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_before = Pt(28)
    r = p.add_run('Documento Tecnico de PDI'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x6B,0x7A,0x8A)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_before = Pt(16)
    for line in [f"Autor: {D['autor']}", f"Unidade: {D['unidade']}", f"Data: {D['data']}",
                 f"Area: {D['area']}", "Status: Entregue (desenvolvido)"]:
        rr = p.add_run(line + '\n'); rr.font.size = Pt(11); rr.font.color.rgb = RGBColor(0x49,0x55,0x60)
    doc.add_page_break()

def add_section(title, blocks):
    hh = doc.add_heading(title, level=2)
    for r in hh.runs: r.font.color.rgb = RGBColor(0x8B,0x1E,0x1E)
    for kind, *rest in blocks:
        if kind == "p":
            doc.add_paragraph(rest[0])
        elif kind in ("ul", "ol"):
            for x in rest[0]:
                doc.add_paragraph(x, style='List Bullet' if kind == "ul" else 'List Number')
        elif kind == "table":
            h, rows = rest
            t = doc.add_table(rows=1, cols=len(h)); t.style = 'Light Grid Accent 1'
            for i, c in enumerate(h): t.rows[0].cells[i].text = str(c)
            for r in rows:
                cells = t.add_row().cells
                for i, c in enumerate(r): cells[i].text = str(c)
        elif kind == "h":
            doc.add_heading(rest[0], level=3)
        elif kind == "note":
            p = doc.add_paragraph(); p.add_run("Nota: ").bold = True; p.add_run(rest[0])
        elif kind == "warn":
            p = doc.add_paragraph(); p.add_run("Atencao: ").bold = True; p.add_run(rest[0])
        elif kind == "code":
            p = doc.add_paragraph(rest[1]); p.style = doc.styles['No Spacing']
            for r in p.runs: r.font.name = 'Consolas'; r.font.size = Pt(9)

cover()
for title, blocks in D['sections']:
    add_section(title, blocks)
out = os.path.join(HERE, "pdi-" + D['slug'] + ".docx")
doc.save(out)
print("DOCX:", out)
