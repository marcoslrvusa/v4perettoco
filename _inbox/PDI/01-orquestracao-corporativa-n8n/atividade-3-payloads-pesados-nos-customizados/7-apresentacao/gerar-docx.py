#!/usr/bin/env python3
"""Generate DOCX from PDI-NOS-CUSTOMIZADOS HTML report."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

def add_cover():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(120)
    run = p.add_run('PDI')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Nós Customizados\ne Expressões Avançadas\nn8n Enterprise')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1F, 0x24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(40)
    run = p.add_run('Relatório Técnico')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x7A, 0x8A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(20)
    for line in ['Autor: Marcos Luciano · PDI Técnico',
                 'Unidade: FV Marketing / V4 Company — Automação & Infraestrutura',
                 'Data: Agosto 2026',
                 'Status: Desenvolvido · NÃO publicado (homologação pendente)',
                 'Versão: 1.0']:
        run = p.add_run(line + '\n')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x49, 0x55, 0x60)

    doc.add_page_break()

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x37, 0x40)
    return h

def add_table(headers, data):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

add_cover()

add_heading_styled('1. Contexto', 1)
doc.add_paragraph(
    'Os pipelines n8n da V4 Company (SDR IA e MarTech) processam payloads cada vez '
    'maiores — listas de leads, sincronizações de CRM, importações em massa. Os nós '
    'Code que fazem esse trabalho foram escritos como "um nó que faz tudo": carregam '
    'o payload inteiro em memória, re-parseiam JSON a cada etapa, percorrem listas com '
    'complexidade O(n²) e repetem a mesma lógica entre dezenas de workflows. Três '
    'problemas estruturais motivam este PDI:'
)
for item in [
    'Payloads pesados estouram a instância — 100k+ itens com map().filter() encadeado '
    'e { ...item } por iteração geram OOM e event loop bloqueado.',
    'Parse e complexidade desperdiçados — JSON.parse dentro de loops e buscas por item '
    '(.find, indexOf) viram O(n²).',
    'Lógica duplicada e sem padrão — o mesmo transform copiado entre workflows, '
    'expressões inline difíceis de manter e nenhuma biblioteca compartilhada.',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('2. Diagnóstico', 1)
add_heading_styled('2.1 Causa raiz', 2)
doc.add_paragraph(
    'Não existe uma camada padrão de transformação no n8n da V4. Cada nó Code reinventa '
    'a roda — e a forma "mais fácil de escrever" é a mais cara de rodar. O resultado é '
    'o sintoma já visto na atividade 1 (ADPLAN: JS timeout de 25 min) e na atividade 2 '
    '(payload pesado sem processamento incremental).'
)

add_heading_styled('3. Solução em 3 frentes', 1)
add_table(
    ['Frente', 'Solução', 'Workflow'],
    [
        ['1. Nós Code JS avançados', 'Streaming com chunking, dedupe O(n) por chave primitiva, cópia mínima e parse 1x.',
         '[CC] NOS - JS Payload Normalizer'],
        ['2. Nós Code Python no n8n', 'Enriquecimento e agregação com stdlib (collections, itertools), sem deps externas.',
         '[CC] NOS - Python Payload Enricher'],
        ['3. Expressões avançadas + biblioteca', 'Referências entre nós, JSONata, condicionais, memoização e lib única 3-lib.',
         '[CC] NOS - Expressions & Memo Playground'],
    ],
)

add_heading_styled('4. Pipeline de processamento', 1)
add_table(
    ['Etapa', 'Responsável', 'Técnica'],
    [
        ['1. Entrada', 'Webhook /nos/js-normalizer', 'Payload bruto (string ou array)'],
        ['2. Parse + Chunk', 'Parse and Chunk (JS)', 'Parse 1x + chunking streaming (default 1000)'],
        ['3. Normalizar', 'Normalize in One Pass (JS)', 'Filtra + dedupe + transforma em O(n)'],
        ['4. Enriquecer', 'Enrich Python (Python)', 'Agregação stdlib (Counter/defaultdict)'],
        ['5. Métricas', 'Return Metrics (JS)', 'durationMs + items/s + dedupe'],
    ],
)
doc.add_paragraph(
    'Escala: o mesmo fluxo que processa 10 itens processa 100k — o chunking streaming '
    'e o dedupe O(n) garantem memória controlada e event loop livre.'
)

add_heading_styled('5. Por que O(n) muda tudo', 1)
doc.add_paragraph(
    'O mesmo payload de 100 mil itens custa bilhões de operações no padrão antigo '
    '(O(n²)) e poucas centenas de milhares no novo (O(n)) — e a memória acompanha. '
    'Chunking streaming mantém o pico de memória constante: o payload inteiro nunca '
    'fica residente.'
)

add_heading_styled('6. Memoização entre execuções', 1)
doc.add_paragraph(
    'Valores estáveis (limiar, taxa, config) são cacheados com $getWorkflowStaticData("global") '
    '— calculados 1x e reutilizados nas execuções seguintes, sem repetir chamadas caras. '
    'Na segunda execução, o cachedAt não muda.'
)

add_heading_styled('7. Biblioteca reutilizável 3-lib', 1)
add_table(
    ['Arquivo', 'Funções'],
    [
        ['payload-lib.js', 'chunk · normalizeStream · dedupe · aggregate · memoizeGlobal · toOutput'],
        ['payload-lib.py', 'chunk · dedupe · aggregate · parse_payload · to_output'],
    ],
)
doc.add_paragraph(
    'Fonte da verdade: nunca editar um nó Code sem atualizar a lib. Os workflows '
    'embutem cópias das funções usadas; a lib é o padrão único.'
)

add_heading_styled('8. Workflows entregues', 1)
add_table(
    ['Workflow', 'Tipo', 'Função'],
    [
        ['[CC] NOS - JS Payload Normalizer', 'Webhook /nos/js-normalizer', 'Parse 1x, chunking streaming, dedupe e normalização O(n)'],
        ['[CC] NOS - Python Payload Enricher', 'Webhook /nos/python-enricher', 'Enriquecimento/agregação com stdlib (set + Counter)'],
        ['[CC] NOS - Expressions & Memo Playground', 'Manual', 'Expressões avançadas + memoização ($getWorkflowStaticData)'],
    ],
)

add_heading_styled('9. Métricas de sucesso', 1)
add_table(
    ['Métrica', 'Antes', 'Meta'],
    [
        ['Payload 100k itens', 'OOM / timeout', 'Streaming, < 2 GB pico'],
        ['Complexidade de normalização', 'O(n²)', 'O(n) padrão'],
        ['Re-parse de JSON por pipeline', 'Múltiplos por etapa', '1x na entrada'],
        ['Código duplicado entre workflows', 'Alto (copiar/colar)', 'Biblioteca única 3-lib'],
        ['Expressões de manutenção difícil', 'Inline, não testáveis', 'Padronizadas + JSONata'],
    ],
)

add_heading_styled('10. Entregas', 1)
add_table(
    ['Pasta', 'Conteúdo', 'Formato'],
    [
        ['1-standards/', '3 padrões: nós Code JS, Python no n8n, expressões avançadas', 'Markdown'],
        ['2-workflows/', '3 workflows validados com n8nac', '.workflow.ts'],
        ['3-lib/', 'Biblioteca reutilizável payload-lib (JS + Python)', 'JS + Python'],
        ['4-retrofit/', 'Como adaptar nós Code existentes (ADPLAN, PRO ANALISES, CC)', 'Markdown'],
        ['5-monitoring/', 'Queries de performance + alertas de duração/dedupe', 'Markdown'],
        ['6-automation/', 'Deploy com gate de aprovação (dry-run)', 'Shell'],
        ['7-apresentacao/', 'Deck HTML + demo + este relatório', 'HTML + Markdown'],
        ['MANUAL-IMPLEMENTACAO.md', 'Manual de implementação completo (6 fases + rollback + checklist)', 'Markdown'],
    ],
)
doc.add_paragraph(
    'Status: desenvolvido · workflows validados e prontos para produção · aguardando '
    'homologação. Nenhum workflow foi publicado no n8n nesta etapa. Implementar '
    'seguindo o MANUAL-IMPLEMENTACAO.md na raiz da atividade.'
)

output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'pdi-nos-customizados-n8n-a3.docx')
doc.save(output_path)
print(f'DOCX saved to {output_path}')