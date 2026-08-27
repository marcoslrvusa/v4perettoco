#!/usr/bin/env python3
"""Generate DOCX from PDI HTML report."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

def add_cover():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(120)
    run = p.add_run('PDI')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Padrão Universal de Tratamento de Erros\nn8n Enterprise')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1F, 0x24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(40)
    run = p.add_run('Relatório Técnico')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x7A, 0x8A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(20)
    for line in ['Autor: Marcos Luciano · PDI Técnico', 'Unidade: FV Marketing / V4 Company — Automação & Infraestrutura',
                 'Data: Julho 2026', 'Status: Homologado', 'Versão: 2.0']:
        run = p.add_run(line + '\n')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x49, 0x55, 0x60)

    doc.add_page_break()

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x37, 0x40)
    return h

add_cover()

add_heading_styled('1. Contexto', 1)
doc.add_paragraph(
    'A operação de Automação & Infraestrutura da V4 Company mantém 10 workflows críticos em produção '
    'no n8n: 7 workflows de SDR IA (atendimento inteligente de leads) e 3 workflows de Command Center '
    '(monitoramento e orquestração da instância).'
)
doc.add_paragraph(
    'Em meados de junho de 2026, identificamos um padrão recorrente: todos os 7 workflows SDR IA '
    'apresentavam falhas silenciosas. Erros aconteciam, os workflows paravam ou degradavam, '
    'e ninguém era notificado. Algumas falhas passavam dias sem detecção, impactando '
    'diretamente o SLA de atendimento aos leads.'
)
doc.add_paragraph(
    'Este relatório documenta o PDI (Padrão Universal de Tratamento de Erros) implementado para resolver '
    'esse problema de raiz — não apenas corrigindo as falhas pontuais, mas estabelecendo um padrão '
    'enterprise reutilizável para todos os workflows da operação.'
)

add_heading_styled('2. Diagnóstico', 1)
add_heading_styled('2.1 Workflows Afetados', 2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Workflow', 'Erro', 'Impacto']
data = [
    ['ADPLAN', 'JS timeout 25min — event loop bloqueado por SplitInBatches ausente', 'P0'],
    ['SIGNOR', 'Task runner desconectado — servidor precisou de tuning', 'P0'],
    ['Genics', 'Redis Cloud inacessível — DNS externo sem fallback', 'P0'],
    ['SOFIA', 'Rate limit Chatwoot — sem wait entre requisições', 'P1'],
    ['PRO ANÁLISES', 'toDateTime undefined — campo opcional sem validação', 'P1'],
    ['Schwalm', 'Null constraint em telefone — insert no Supabase sem validação', 'P1'],
    ['V4 INTERNO', '404 no Ekyte — board ID mudou sem atualizar o workflow', 'P1'],
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

add_heading_styled('2.2 Causa Raiz', 2)
doc.add_paragraph(
    'Ausência de padronização: cada workflow tratava erro do seu jeito improvisado — ou '
    'simplesmente não tratava. Não havia taxonomia de erros, política de retry, notificação, '
    'persistência ou proteção contra cascata. O resultado era um ecossistema frágil onde uma falha '
    'em um nó poderia derrubar todo o fluxo sem deixar rastro.'
)

add_heading_styled('3. Arquitetura da Solução', 1)
doc.add_paragraph(
    'Três camadas obrigatórias e complementares. Não são alternativas — um workflow enterprise precisa das três.'
)
add_heading_styled('Camada 1: Node-Level · Auto-cura local', 2)
doc.add_paragraph(
    'retryOnFail + continueErrorOutput em cada nó falível. Captura ~73% das falhas transientes '
    '(429, 503, timeout) no próprio nó — sem overhead de workflow externo.'
)
add_heading_styled('Camada 2: Error Handler Central · Rede de segurança', 2)
doc.add_paragraph(
    'Workflow único que captura tudo que escapar da Camada 1. Classifica severidade por taxonomia, '
    'notifica Slack em < 1 min, persiste na Dead Letter Queue. 10 nós, validado com n8nac.'
)
add_heading_styled('Camada 3: Circuit Breaker + Dead Letter Queue · Proteção sistêmica', 2)
doc.add_paragraph(
    'Circuit breaker evita sobrecarga em APIs fragilizadas (abre após 5 falhas consecutivas, '
    'cooldown de 5 min, recovery automático). DLQ preserva payload completo para replay e auditoria.'
)

add_heading_styled('4. Taxonomia de Erros', 1)
doc.add_paragraph(
    'Todo erro é classificado no momento da captura em exatamente uma das 8 classes abaixo. '
    'Sem ambigüidade, sem exceção.'
)
table = doc.add_table(rows=9, cols=5)
table.style = 'Light Grid Accent 1'
headers = ['Classe', 'Severidade', 'Retentável?', 'Exemplo', 'Ação']
data = [
    ['server_error', 'critical', 'Sim (3x)', '502, 503', 'Retry + DLQ + Slack #incidents'],
    ['rate_limit', 'warning', 'Sim (c/ backoff)', '429', 'Backoff + jitter + DLQ'],
    ['timeout', 'critical', 'Sim (3x)', 'ETIMEDOUT', 'Retry + DLQ + Slack'],
    ['client_error', 'warning', 'Nunca', '400, 401, 403', 'DLQ + notificar equipe'],
    ['data_validation', 'warning', 'Nunca', 'null constraint', 'DLQ + notificar'],
    ['network_dns', 'critical', 'Sim (3x)', 'ENOTFOUND', 'Retry + DLQ + Slack'],
    ['runtime', 'critical', 'Nunca', 'Task runner off', 'Alarmar imediatamente'],
    ['resource_exhaustion', 'critical', 'Nunca', 'OOM', 'Alarmar imediatamente'],
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

add_heading_styled('5. Envelope Padrão de Erro', 1)
doc.add_paragraph(
    'Todo workflow que captura um erro DEVE produzir um envelope JSON com: severity, errorClass, '
    'correlationId (estável entre retries), workflowName, workflowId, failedNode, errorMessage, '
    'executionUrl e timestamp.'
)

p = doc.add_paragraph()
run = p.add_run(
    '{\n'
    '  "severity": "critical",\n'
    '  "errorClass": "server_error",\n'
    '  "correlationId": "abc123-67890",\n'
    '  "workflowName": "CRM Sync",\n'
    '  "workflowId": "abc123",\n'
    '  "failedNode": "HTTP Request — Attio",\n'
    '  "errorMessage": "Request failed with status code 502",\n'
    '  "executionUrl": "https://n8n.fvmarketing.com.br/workflow/abc123/executions/67890",\n'
    '  "timestamp": "2026-07-08T10:30:00.000Z"\n'
    '}'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading_styled('6. Retry Matrix', 1)
doc.add_paragraph('Cada tipo de nó tem sua própria política de retry. Backoff exponencial com jitter para 429 e 5xx.')
table = doc.add_table(rows=10, cols=5)
table.style = 'Light Grid Accent 1'
headers = ['Tipo de Nó', 'retryOnFail', 'maxTries', 'waitBetweenTries', 'Observação']
data = [
    ['HTTP Request', 'Sim', '3', '5000 ms', 'Sempre configurar'],
    ['Supabase', 'Sim', '2', '3000 ms', 'Falhas de rede raras'],
    ['Slack', 'Sim', '3', '5000 ms', 'Rate limit 429 frequente'],
    ['WhatsApp / API externa', 'Sim', '3', '5000 ms', 'Instabilidade comum'],
    ['Google Sheets / Drive', 'Sim', '2', '5000 ms', '429 em pico'],
    ['Redis', 'Sim', '2', '2000 ms', 'Conexão local'],
    ['Postgres', 'Sim', '2', '3000 ms', 'Timeout de query'],
    ['n8n API', 'Sim', '3', '5000 ms', 'Rate limit interno'],
    ['Code node', 'Não', '—', '—', 'NUNCA retentar — erro de lógica'],
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

add_heading_styled('7. Error Handler Central', 1)
doc.add_paragraph(
    'O workflow central que orquestra todo o tratamento de erro. Recebe erros de todos os workflows, '
    'classifica, notifica e persiste.'
)
doc.add_paragraph(
    'Pipeline: Error Trigger → Code (Parse & Classify) → IF Severity Critical? '
    '→ Sim: Slack #incidents + Supabase INSERT DLQ + HTTP PATCH Circuit Breaker '
    '| Não: Slack #alerts + Supabase INSERT DLQ'
)

add_heading_styled('8. Circuit Breaker & Dead Letter Queue', 1)
add_heading_styled('8.1 Circuit Breaker', 2)
doc.add_paragraph('Protege APIs contra cascata de falhas. O estado é isolado por workflow.')
doc.add_paragraph(
    'Parâmetros: FAILURE_THRESHOLD = 5 falhas consecutivas, COOLDOWN_MS = 300.000 (5 min), '
    'Recovery automático via PATCH n8n API, Monitor a cada 5 min.'
)
doc.add_paragraph('Ciclo: closed → (5 falhas) → open → (5 min) → half-open → (sucesso) → closed')

add_heading_styled('8.2 Dead Letter Queue', 2)
doc.add_paragraph(
    'Persistência permanente de todas as falhas não-recuperáveis. Diferente do prune de execuções '
    'do n8n (que apaga dados antigos), a DLQ preserva para auditoria.'
)
doc.add_paragraph(
    'Campos: correlation_id (TEXT), error_class (TEXT com CHECK), severity (TEXT com CHECK), '
    'payload (JSONB), status (TEXT: pending → investigating → resolved)'
)

add_heading_styled('9. Plano de Retrofit', 1)
doc.add_paragraph(
    '5 workflows SDR IA já tiveram correções aplicadas nos JSONs locais em 22/06/2026. '
    'Estão pendentes de push para o n8n. O retrofit completo está estimado em 4 dias.'
)
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['Workflow', 'Erro', 'Correção', 'Esforço']
data = [
    ['ADPLAN', 'JS timeout 25min', 'SplitInBatches + Code otimizado', '20 min'],
    ['SOFIA', 'Rate limit Chatwoot', 'Wait 500ms + Error Workflow', '10 min'],
    ['PRO ANÁLISES', 'toDateTime undefined', 'try/catch toDateTime', '10 min'],
    ['Schwalm', 'Null constraint telefone', 'Validação + IF', '10 min'],
    ['V4 INTERNO', '404 Ekyte board ID', 'Board ID map + IF', '10 min'],
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

add_heading_styled('10. Anti-Patterns', 1)
doc.add_paragraph('Erros reais observados nos workflows SDR IA que o padrão veio eliminar:')
anti_patterns = [
    'onError sem main[1] — Erro descartado silenciosamente. Sempre usar continueErrorOutput + segundo output conectado ao handler.',
    'Retentar 4xx — Queima API credits à toa. Erro do cliente (400, 401, 403, 404) não vai mudar na próxima tentativa. Só retentar 5xx / timeout / DNS.',
    'Error workflow no mesmo canal — Se o workflow de erro usa o mesmo Slack que o principal, uma falha no Slack cria recursão infinita. Usar canais separados.',
    '200 no erro — Se a resposta é 200 mas houve falha interna, quem chamou nunca sabe. Propagar status HTTP correto.',
    'Não publicar o error handler — Código antigo rodando mesmo depois de "corrigido". Sempre publish (Shift+P) após configurar ou alterar.',
    'Payload sem correlationId — Sem um ID estável entre retries, não é possível ligar uma falha à sua causa original.',
]
for ap in anti_patterns:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(ap)
    run.font.size = Pt(10)

add_heading_styled('11. Métricas de Sucesso', 1)
metrics = [
    ('Detecção de falha (antes)', 'Dias', 'bad'),
    ('Detecção de falha (meta)', '< 1 min', 'good'),
    ('Auto-cura retry (antes)', '0%', 'bad'),
    ('Auto-cura retry (meta)', '> 70%', 'good'),
    ('Circuitos abertos sem alerta (antes)', '100%', 'bad'),
    ('Circuitos abertos sem alerta (meta)', '0%', 'good'),
    ('Workflows c/ error handling (antes)', '0', 'bad'),
    ('Workflows c/ error handling (meta)', '100%', 'good'),
]
for label, val, _ in metrics:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    p.add_run(val)

add_heading_styled('12. Entregas do PDI', 1)
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['Pasta', 'Conteúdo', 'Formato']
data = [
    ['1-standards/', 'Documentação do padrão, taxonomia, envelope de erro, retry matrix (609 linhas)', 'Markdown'],
    ['2-workflows/', 'Error Handler Central + Circuit Breaker Monitor, validados com n8nac', '.workflow.ts'],
    ['3-supabase/', 'Schema v2.1 (4 tabelas + 4 views) + guia de migração', 'SQL + Markdown'],
    ['4-retrofit/', 'Planos de retrofit para SDR IA (7 workflows) e Command Center (3 workflows)', 'Markdown'],
    ['5-monitoring/', 'Consultas SQL de dashboard e regras de alerta', 'Markdown'],
    ['6-automation/', 'Scripts de deploy automatizado (validação + push + migração)', 'Shell'],
    ['7-apresentacao/', 'Deck interativo HTML (14 slides) + script de demonstração + este relatório', 'HTML + Markdown'],
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

output_path = '/home/marcos/Desktop/AI/v4perettoco-main/PDI/7-apresentacao/pdi-orquestracao-corporativa-n8n-a1.docx'
doc.save(output_path)
print(f'DOCX saved to {output_path}')
