#!/usr/bin/env python3
"""Generate DOCX from PDI-MARTECH HTML report."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

def add_cover():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(120)
    run = p.add_run('PDI')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Resiliência MarTech\nn8n Enterprise')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1F, 0x24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(40)
    run = p.add_run('Relatório Técnico')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x7A, 0x8A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(20)
    for line in ['Autor: Marcos Luciano · PDI Técnico',
                 'Unidade: FV Marketing / V4 Company — Automação & Infraestrutura',
                 'Data: Agosto 2026',
                 'Status: Desenvolvido · NÃO publicado (homologação pendente)',
                 'Versão: 1.0']:
        run = p.add_run(line + '\n')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x49, 0x55, 0x60)

    doc.add_page_break()

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x37, 0x40)
    return h

def add_table(headers, data):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

add_cover()

add_heading_styled('1. Contexto', 1)
doc.add_paragraph(
    'A operação MarTech da V4 Company processa requisições de integração com CRMs '
    '(Kommo, HubSpot, RD Station, entre outros) em workflows n8n. Três problemas '
    'estruturais motivam este PDI:'
)
for item in [
    'Workflows síncronos frágeis — requisição processada dentro do webhook; em picos '
    '(campanha, importação), a instância trava e o cliente não recebe confirmação.',
    'Payloads pesados sem checkpoint — processamento de 10k+ itens de uma vez; '
    'um timeout no meio perde tudo e o job recomeça do zero.',
    'Sincronização com CRM sem observabilidade — falha não registrada; divergência '
    'entre o esperado e o syncado só é descoberta quando o cliente reclama.',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('2. Diagnóstico', 1)
add_heading_styled('2.1 Causa raiz', 2)
doc.add_paragraph(
    'Trabalho de integração tratado como transação síncrona. Sem fila, um pico vira '
    'travamento; sem checkpoint, um timeout vira retrabalho total; sem trilha de sync, '
    'a falha fica invisível até o impacto no cliente.'
)

add_heading_styled('3. Solução em 3 frentes', 1)
add_table(
    ['Frente', 'Solução', 'Workflow'],
    [
        ['1. Fila + concorrência', 'Webhook enfileira e responde ACK 202 imediato; worker consome respeitando o semáforo de slots por fila.',
         '[CC] MT Queue Gateway + [CC] MT - Queue Worker'],
        ['2. Payload pesado', 'Processamento em chunks (JS) + enriquecimento (Python) com checkpoint retomável.',
         '[CC] MT - Heavy Payload Processor'],
        ['3. Observabilidade CRM', 'Auditoria de sync, health por objeto e detecção de divergência (drift) antes do cliente.',
         '[CC] MT - CRM Sync Observabilidade'],
    ],
)

add_heading_styled('4. Pipeline de processamento', 1)
add_table(
    ['Etapa', 'Responsável', 'Persistência'],
    [
        ['1. Enfileirar', 'MT Queue Gateway', 'mt_jobs (queued)'],
        ['2. Consumir', 'MT - Queue Worker (15s)', 'slot mt_concurrency → running'],
        ['3. Processar', 'MT - Heavy Payload Processor', 'checkpoint mt_job_progress'],
        ['4. Observar', 'MT - CRM Sync Observabilidade', 'mt_sync_log + health + delta'],
    ],
)
doc.add_paragraph(
    'Retomada: falha no chunk 4/10 → o worker re-enfileira com backoff exponencial '
    '(30s → 1m → 2m, máx. 3) e o processamento retoma do chunk 4, não do zero.'
)

add_heading_styled('5. Detecção de divergência (drift)', 1)
doc.add_paragraph(
    'Todo sync envia um envelope padrão (expected × confirmed). O Observabilidade '
    'calcula o drift % e decide: divergência acima da tolerância (default 5%) gera '
    'mt_sync_delta aberto + alerta no dashboard — antes do impacto chegar ao cliente.'
)

add_heading_styled('6. Schema Supabase v3.0', 1)
add_table(
    ['Tabela', 'Uso'],
    [
        ['mt_jobs', 'Fila assíncrona (queued/running/done/failed, backoff, heartbeat)'],
        ['mt_concurrency', 'Semáforo distribuído — limite por fila'],
        ['mt_job_progress', 'Checkpoint de chunk de payload pesado'],
        ['mt_sync_log', 'Auditoria de cada sync de CRM'],
        ['mt_crm_health', 'Saúde agregada por object + direction'],
        ['mt_sync_delta', 'Divergências abertas (drift)'],
    ],
)
doc.add_paragraph(
    'Aditivo: não altera tabelas/views v2.x (error_*) da atividade 1. Convive lado a lado.'
)
doc.add_paragraph(
    'Views: vw_mt_queue_backlog · vw_mt_slots · vw_mt_sync_summary_24h · '
    'vw_mt_drift_abertos · vw_mt_crm_health'
)

add_heading_styled('7. Workflows entregues', 1)
add_table(
    ['Workflow', 'Tipo', 'Função'],
    [
        ['[CC] MT Queue Gateway', 'Webhook', 'Enfileira + ACK 202 + idempotência por job_key'],
        ['[CC] MT - Queue Worker', 'Schedule 15s', 'Consome slots, delega, faz retry com backoff'],
        ['[CC] MT - Heavy Payload Processor', 'Sub-workflow', 'Chunks JS + Python, checkpoint e retomada'],
        ['[CC] MT - CRM Sync Observabilidade', 'Webhook /mt/crm-sync', 'Log, health e drift'],
    ],
)

add_heading_styled('8. Métricas de sucesso', 1)
add_table(
    ['Métrica', 'Antes', 'Meta'],
    [
        ['ACK de requisição MarTech', 'Trava no webhook', '< 2s'],
        ['Job pesado retomável', 'Não', 'Sim'],
        ['Falha de sync detectada', 'Quando o cliente reclama', '< 15 min'],
        ['Divergência (drift) visível', 'Não', 'Dashboard em tempo real'],
        ['Concorrência controlada', 'Manual/improvisada', 'Semáforo mt_concurrency'],
    ],
)

add_heading_styled('9. Entregas', 1)
add_table(
    ['Pasta', 'Conteúdo', 'Formato'],
    [
        ['1-standards/', '3 padrões: filas/concorrência, payloads pesados, observabilidade CRM', 'Markdown'],
        ['2-workflows/', '4 workflows validados com n8nac', '.workflow.ts'],
        ['3-supabase/', 'Schema v3.0 (6 tabelas + 5 views) + guia de migração', 'SQL + Markdown'],
        ['4-retrofit/', 'Como adaptar workflows MarTech existentes', 'Markdown'],
        ['5-monitoring/', 'Consultas SQL de dashboard + alertas', 'Markdown'],
        ['6-automation/', 'Scripts de migração e deploy (com gate de aprovação)', 'Shell'],
        ['7-apresentacao/', 'Deck HTML + demo + este relatório', 'HTML + Markdown'],
    ],
)
doc.add_paragraph(
    'Status: desenvolvido · workflows e schema prontos para produção · aguardando '
    'homologação. Nenhum workflow foi publicado no n8n nesta etapa.'
)

output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'pdi-orquestracao-martech-n8n-a2.docx')
doc.save(output_path)
print(f'DOCX saved to {output_path}')