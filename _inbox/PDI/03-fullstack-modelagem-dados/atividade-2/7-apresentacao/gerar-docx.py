#!/usr/bin/env python3
"""Gera o DOCX desta atividade a partir de report.json (padrão PDI)."""
import json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
with open(os.path.join(HERE, "report.json"), encoding="utf-8") as f:
    A = json.load(f)

doc = Document()
st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3); s.right_margin = Cm(3)

def cover():
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_before = Pt(120)
    r = p.add_run('PDI'); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x8B,0x45,0x13)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(A['title']); r.bold = True; r.font.size = Pt(24)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_before = Pt(30)
    r = p.add_run('Relatório Técnico'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x6B,0x7A,0x8A)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_before = Pt(16)
    for line in [f"Autor: {A['autor']}", f"Unidade: {A['unidade']}", f"Data: {A['data']}",
                 f"Área: {A['area']}", "Status: Entregue (desenvolvido)"]:
        rr = p.add_run(line + '\n'); rr.font.size = Pt(11); rr.font.color.rgb = RGBColor(0x49,0x55,0x60)
    doc.add_page_break()

def h(t, lvl=1):
    x = doc.add_heading(t, level=lvl)
    for run in x.runs:
        run.font.color.rgb = RGBColor(0x8B,0x45,0x13) if lvl==1 else RGBColor(0x2E,0x37,0x40)
    return x

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

cover()
h('1. Contexto e Problema'); doc.add_paragraph(A['problem'])
h('2. Diagnóstico'); bullets(A['diagnosis'])
h('3. Arquitetura da Solução'); bullets(A['architecture'])
h('4. Entregas'); bullets(A['deliverables'])
h('5. Métricas de Sucesso')
t = doc.add_table(rows=1, cols=3); t.style = 'Light Grid Accent 1'
for i,c in enumerate(['Métrica','Antes','Depois']): t.rows[0].cells[i].text = c
for l,b,aft in A['metrics']:
    row = t.add_row().cells; row[0].text=l; row[1].text=b; row[2].text=aft
h('6. Próximos Passos'); bullets(A['next'])

out = os.path.join(HERE, A['docx'])
doc.save(out)
print('DOCX:', out)
